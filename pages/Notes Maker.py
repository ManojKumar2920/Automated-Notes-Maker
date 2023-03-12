import streamlit as st
import speech_recognition as sr
import docx
from pydub import AudioSegment
from matplotlib import image
import os
import io



FILE_DIR = os.path.dirname(os.path.abspath(__file__))
PARENT_DIR = os.path.join(FILE_DIR, os.pardir)
dir_of_interest = os.path.join(PARENT_DIR, "resources")
IMAGE_PATH = os.path.join(dir_of_interest, "Notes_maker.jpeg")

img = image.imread(IMAGE_PATH)

st.image(img,width=400)

def convert_to_wav(audio):
    if audio.split('.')[-1]=='wav':
        return audio
    else:
        new_audio='.'.join(audio.split('.')[:-1])+'.wav'
        Audio=AudioSegment.from_file(audio)
        Audio.export(new_audio,format='wav')
        return new_audio

def is_audio_clear(audio):
    sound = AudioSegment.from_file(audio)
    if sound.dBFS > -50:
        return True
    else:
        return False

def audio_to_text(audio):
    r = sr.Recognizer()
    with sr.AudioFile(audio) as source:
        audio_text = r.record(source)
    return r.recognize_google(audio_text)

    
def create_document(filename, text, filepath):
    if os.path.isfile(filename):
        doc = docx.Document(filename)
    else:
        doc = docx.Document()
        doc.add_heading('Transcribed Notes',0)
    doc.add_paragraph(text)
    doc.save(filepath)

def convert_to_docx(audio, filepath):
    if not is_audio_clear(audio):
        st.warning("Audio quality is too low. Please upload a clearer audio file.")
        return
    wav_file=convert_to_wav(audio)
    
    # Add a loading spinner while the file is being converted
    with st.spinner("Converting file..."):
        text=audio_to_text(wav_file)
        filename=os.path.splitext(audio)[0]
        create_document(filename, text, filepath)
    
    st.success("File successfully converted to text")
    
    # Print transcribed text
    st.subheader("Transcribed Text")
    st.write(text)
    
    # Download transcribed notes
    with io.BytesIO() as f:
        create_document(filename, text, f)
        f.seek(0)
        st.download_button(label="Download", data=f.getvalue(), file_name="Transcribed Notes.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.write("Thank you for using our Audio to Text Converter!")


st.markdown("""
    <style>
        body {
            background-color: #f0f0f0;
        }
        h1, h2, h3, h4, h5, h6 {
            font-family: 'Roboto', sans-serif;
            color: #333333;
        }
        .btn-primary {
            background-color: #007bff;
            border-color: #007bff;
            font-family: 'Roboto', sans-serif;
            font-weight: bold;
        }
        .btn-primary:hover {
            background-color: #0069d9;
            border-color: #0062cc;
        }
    </style>
""", unsafe_allow_html=True)

st.title("Audio to Text Converter")

uploaded_file = st.file_uploader("Upload an audio file", type=["wav"])

if uploaded_file is not None:
    file_ext = uploaded_file.name.split(".")[-1]
    if file_ext in ["wav"]:
        with open("temp." + file_ext, "wb") as f:
            f.write(uploaded_file.getbuffer())
        st.success("File uploaded
